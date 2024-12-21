import streamlit as st
import requests
import time
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from moviepy.editor import AudioFileClip
import tempfile
import os
from dontenv import load_dotenv

# Set your Hugging Face API token here (ensure this is kept secure)
hf_token = os.getenv('href')

# Custom CSS
st.markdown(
    """
    <style>
    .title {
        text-align: center;
        font-size: 2.5em;
        color: #4CAF50;
    }
    .header {
        text-align: left;
        font-size: 1.5em;
        color: #4CAF50;
    }
    .download-btn {
        background-color: #4CAF50; 
        color: white; 
        padding: 10px 20px; 
        border: none; 
        border-radius: 5px;
        cursor: pointer;
    }
    .download-btn:hover {
        background-color: #45a049;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Function to convert audio file to WAV format using moviepy
def convert_to_wav(audio_file):
    try:
        # Create a temporary file to save the uploaded audio
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as temp_file:
            # Save the uploaded audio to the temporary file
            temp_file.write(audio_file.read())
            temp_file_path = temp_file.name  # Get the file path

        # Load audio file and convert to WAV format
        audio = AudioFileClip(temp_file_path)
        wav_file_path = "converted_audio.wav"
        # Export the audio to WAV format
        audio.write_audiofile(wav_file_path, codec='pcm_s16le')
        
        # Clean up: remove the temporary file
        os.remove(temp_file_path)
        
        return wav_file_path  # Return the path to the converted file
    except Exception as e:
        print(f"Error converting audio: {e}")  # Print error details
        return None

# Function to transcribe audio
def transcribe_audio(audio_file):
    url = "https://api-inference.huggingface.co/models/facebook/wav2vec2-base-960h"
    headers = {"Authorization": f"Bearer {hf_token}"}
    
    # Convert audio to WAV format
    audio_data = convert_to_wav(audio_file)
    
    if audio_data is None:
        st.write("Audio conversion failed.")  # Feedback on failure
        return None
    
    # Transcription request loop
    while True:
        try:
            with open(audio_data, "rb") as file:  # Open the converted WAV file
                response = requests.post(url, headers=headers, data=file.read())
        except Exception as e:
            st.write(f"Error during request: {e}")  # Print request error
            return None
        
        if response.status_code == 503:
            st.write("Model is loading... retrying in 30 seconds.")
            time.sleep(30)  # Wait for 30 seconds and retry
        elif response.status_code == 200:
            transcription = response.json().get("text", "Transcription failed")
            return transcription
        else:
            st.write(f"Error in transcription: {response.text}")
            return None

# Function to summarize the transcription text
def summarize_text(transcription):
    url = "https://api-inference.huggingface.co/models/facebook/bart-large-cnn"
    headers = {"Authorization": f"Bearer {hf_token}"}
    
    # Set up data with parameters for summary
    data = {
        "inputs": transcription,
        "parameters": {
            "max_length": 50,  # Limit summary length
            "min_length": 10,  # Ensure summary is not too short
            "do_sample": False  # Use deterministic output
        }
    }

    # Summarization request
    response = requests.post(url, headers=headers, json=data)

    if response.status_code == 200:
        summary = response.json()[0].get("summary_text", "Summarization failed")
        return summary
    else:
        st.write(f"Error in summarization: {response.text}")
        return None

# Function to generate a .doc file with MOM structure
def create_doc(transcription, summary):
    date = time.strftime("%Y-%m-%d")  # Get the current date
    doc = Document()
    doc.add_heading('Minutes of Meeting (MOM)', level=1)
    doc.add_heading('Agenda:', level=2)
    doc.add_paragraph('Audio Transcription and Summarization')
    doc.add_heading('Date:', level=2)
    doc.add_paragraph(date)
    
    doc.add_paragraph("---")
    doc.add_heading('Transcription:', level=2)
    doc.add_paragraph(transcription)
    
    doc.add_paragraph("---")
    doc.add_heading('Summary:', level=2)
    doc.add_paragraph(summary)
    
    doc.add_paragraph("---")
    doc.add_heading('Action Items:', level=2)
    doc.add_paragraph("1. Review the transcription and summary.")
    doc.add_paragraph("2. Identify key points for follow-up.")
    doc.add_paragraph("3. Distribute the MOM to participants.")
    
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

# Function to generate a .pdf file with MOM structure
def create_pdf(transcription, summary):
    date = time.strftime("%Y-%m-%d")  # Get the current date
    pdf_stream = BytesIO()
    c = canvas.Canvas(pdf_stream, pagesize=letter)
    
    c.drawString(100, 750, "Minutes of Meeting (MOM)")
    c.drawString(100, 730, "Agenda: Audio Transcription and Summarization")
    c.drawString(100, 710, f"Date: {date}")
    
    c.drawString(100, 690, "---")
    c.drawString(100, 670, "Transcription:")
    c.drawString(100, 650, transcription)
    
    c.drawString(100, 630, "---")
    c.drawString(100, 610, "Summary:")
    c.drawString(100, 590, summary)
    
    c.drawString(100, 570, "---")
    c.drawString(100, 550, "Action Items:")
    c.drawString(100, 530, "1. Review the transcription and summary.")
    c.drawString(100, 510, "2. Identify key points for follow-up.")
    c.drawString(100, 490, "3. Distribute the MOM to participants.")
    
    c.save()
    pdf_stream.seek(0)
    return pdf_stream

# Streamlit App
def main():
    # Initialize session state for history if it doesn't exist
    if 'history' not in st.session_state:
        st.session_state.history = []

    st.sidebar.title("Legallify")
    selection = st.sidebar.radio("Navigate", 
                                  ["🏠 Home", "📤 Upload File", "🗂️ History"])

    if selection == "🏠 Home":
        st.markdown('<h1 class="title">Welcome to Legallify</h1>', unsafe_allow_html=True)
        st.write("Legallify uses state-of-the-art machine learning models to transcribe audio files and generate summaries.")
        st.write("Upload your audio files in WAV, MP3, or MP4 formats to get started!")
        st.write("### Model Descriptions:")
        st.write("- **Transcription Model:** Utilizes the `facebook/wav2vec2-base-960h` model to convert audio into text.")
        st.write("- **Summarization Model:** Uses the `facebook/bart-large-cnn` model to condense text into a summary.")

    elif selection == "📤 Upload File":
        st.markdown('<h1 class="title">Audio to Text Transcription & Summarization</h1>', unsafe_allow_html=True)

        # File uploader for audio files
        audio_file = st.file_uploader("🚀 Tap to Browse and Upload Your Audio File!", type=["wav", "mp3", "mp4"])

        if audio_file is not None:
            st.write("Processing audio file...")
            with st.spinner("Transcribing... Please wait."):
                # Transcription
                transcription = transcribe_audio(audio_file)
            
            if transcription:
                # Append to history with initial status
                entry = {
                    "file_name": audio_file.name,
                    "transcribed": True,
                    "summary": None  # Summary starts as None
                }
                st.session_state.history.append(entry)

                # Summarization
                st.write("\nSummarizing the transcription...")
                with st.spinner("Summarizing... Please wait."):
                    summary = summarize_text(transcription)

                if summary:
                    # Update history with summary status
                    st.session_state.history[-1]["summary"] = summary  # Save summary to history

                    # Download options
                    st.write("Download options:")
                    download_format = st.selectbox("Choose format:", ["Word Document (.doc)", "PDF (.pdf)"])

                    if download_format == "Word Document (.doc)":
                        doc_stream = create_doc(transcription, summary)
                        st.download_button(label="📥 Download MOM (Word Document)", data=doc_stream, 
                                           file_name=f"mom_{time.strftime('%Y-%m-%d')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                    elif download_format == "PDF (.pdf)":
                        pdf_stream = create_pdf(transcription, summary)
                        st.download_button(label="📥 Download MOM (PDF)", data=pdf_stream, 
                                           file_name=f"mom_{time.strftime('%Y-%m-%d')}.pdf", mime="application/pdf")

    elif selection == "🗂️ History":
        st.markdown('<h1 class="title">Transcription & Summarization History</h1>', unsafe_allow_html=True)

        if len(st.session_state.history) > 0:
            for idx, entry in enumerate(st.session_state.history):
                st.write(f"**{idx + 1}. File Name:** {entry['file_name']}")
                st.write(f"**Transcription Status:** {'Completed' if entry['transcribed'] else 'Pending'}")
                st.write(f"**Summary Status:** {'Completed' if entry['summary'] else 'Pending'}")
                st.write("---")
        else:
            st.write("No history available.")

if __name__ == "__main__":
    main()
